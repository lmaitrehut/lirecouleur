import re
import pyphen
from docx import Document
from docx.shared import RGBColor

try:
    dic = pyphen.Pyphen(lang="fr")
except Exception:
    try:
        dic = pyphen.Pyphen(lang="fr_FR")
    except Exception:
        dic = None

VERT = RGBColor(0, 150, 0)
BLEU = RGBColor(0, 32, 255)
ROUGE = RGBColor(220, 0, 0)
COULEURS = [BLEU, ROUGE]


def _decouper_fallback(mot):
    pattern = r"(?i)[^aeiouyàâéèêëîïôûùüç]*[aeiouyàâéèêëîïôûùüç]+(?:(?![aeiouyàâéèêëîïôûùüç])[^aeiouyàâéèêëîïôûùüç])?"
    syllabes = re.findall(pattern, mot)
    if syllabes and "".join(syllabes) == mot:
        return syllabes
    return [mot]


def _syllabes(mot):
    match = re.match(r"^([^\wàâéèêëîïôûùüç]*)(.*?)([^\wàâéèêëîïôûùüç]*)$", mot, re.IGNORECASE)
    if not match:
        return [mot]
    prefixe, mot_epure, suffixe = match.groups()
    if not mot_epure:
        return [mot]
    syllabes = []
    if dic:
        decoupe = dic.inserted(mot_epure)
        if decoupe and "-" in decoupe:
            syllabes = decoupe.split("-")
    if not syllabes:
        syllabes = _decouper_fallback(mot_epure)
    if prefixe:
        syllabes[0] = prefixe + syllabes[0]
    if suffixe:
        syllabes[-1] = syllabes[-1] + suffixe
    return syllabes


def process_docx(input_path, output_path):
    doc = Document(input_path)

    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        mots = p.text.split()
        p.clear()  # supprime proprement tous les runs existants
        for idx_mot, mot in enumerate(mots):
            if not mot:
                continue
            syllabes = _syllabes(mot)
            premiere_lettre = True  # premiere lettre du mot -> verte
            for idx_syl, syl in enumerate(syllabes):
                couleur = COULEURS[idx_syl % len(COULEURS)]
                # Gestion de la toute premiere lettre (verte + majuscule)
                if premiere_lettre:
                    # trouver la premiere lettre dans la syllabe (ignore ponctuation)
                    pos = next((i for i, c in enumerate(syl) if c.isalpha()), None)
                    if pos is not None:
                        # run pour la premiere lettre (vert + majuscule)
                        r = p.add_run(syl[pos].upper())
                        r.font.color.rgb = VERT
                        # reste avant/du milieu en couleur de syllabe
                        if pos > 0:
                            rpre = p.add_run(syl[:pos])
                            rpre.font.color.rgb = couleur
                        # reste apres la premiere lettre en couleur de syllabe
                        if pos + 1 < len(syl):
                            rapp = p.add_run(syl[pos + 1:])
                            rapp.font.color.rgb = couleur
                        premiere_lettre = False
                    else:
                        # ponctuation seule dans la syllabe (ex: "'")
                        r = p.add_run(syl)
                        r.font.color.rgb = couleur
                else:
                    # autres syllabes : un seul run colore
                    r = p.add_run(syl)
                    if any(c.isalnum() for c in syl):
                        r.font.color.rgb = couleur
            if idx_mot < len(mots) - 1:
                p.add_run("  ")

    doc.save(output_path)
    return output_path
